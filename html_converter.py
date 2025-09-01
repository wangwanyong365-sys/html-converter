import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class HTMLConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("HTML 转换器")
        # 调整窗口大小并设置最小尺寸
        self.root.geometry("550x450") # 增加高度以容纳新控件
        self.root.minsize(500, 500)

        # 设置主题和自定义样式
        self.style = ttk.Style(self.root)
        self.style.theme_use("clam")
    
        # --- 自定义按钮样式 ---
        # 获取主题的默认背景色
        bg_color = self.root.cget('bg')
        
        self.style.configure("Accent.TButton", 
                             font=("微软雅黑", 12, "bold"), 
                             padding=10,
                             background=bg_color,
                             foreground="black",
                             borderwidth=1,
                             relief="solid")
        
        # 设置鼠标悬浮时的样式
        self.style.map("Accent.TButton",
            background=[('active', '#e0e0e0')], # 鼠标悬浮时变色
            relief=[('pressed', 'sunken')]
        )
    
        self.setup_ui()
    
    def setup_ui(self):
        # ---- 主框架 ----
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # 使网格布局自适应窗口大小
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(2, weight=1)
    
        # ---- 标题 ----
        title_label = ttk.Label(main_frame, text="HTML 转换工具", font=("微软雅黑", 20, "bold"))
        title_label.grid(row=0, column=0, pady=(0, 20), sticky=tk.N)
    
        # ---- 文件选择区域 ----
        files_frame = ttk.LabelFrame(main_frame, text="源文件选择", padding="10")
        files_frame.grid(row=1, column=0, pady=10, sticky=(tk.W, tk.E, tk.N))
        files_frame.columnconfigure(0, weight=1)
    
        self.html_entry = ttk.Entry(files_frame, font=("Arial", 10))
        self.html_entry.grid(row=0, column=0, ipady=4, sticky=(tk.W, tk.E))
        self.browse_btn = ttk.Button(files_frame, text="浏览...", command=self.browse_html)
        self.browse_btn.grid(row=0, column=1, padx=(10, 0))
    
        # ---- 输出设置 ----
        options_frame = ttk.LabelFrame(main_frame, text="输出设置", padding="10")
        options_frame.grid(row=2, column=0, pady=10, sticky=(tk.W, tk.E, tk.N))
        options_frame.columnconfigure(1, weight=1)
        
        # --- 新增: 输出目录选择 ---
        ttk.Label(options_frame, text="输出目录:").grid(row=0, column=0, padx=(0, 10), pady=(0, 5), sticky=tk.W)
        self.output_dir_entry = ttk.Entry(options_frame, font=("Arial", 10))
        self.output_dir_entry.grid(row=0, column=1, pady=(0, 5), sticky=(tk.W, tk.E))
        self.browse_output_btn = ttk.Button(options_frame, text="选择...", command=self.browse_output_dir)
        self.browse_output_btn.grid(row=0, column=2, padx=(10, 0), pady=(0, 5))

        # --- 修改: 调整行号 ---
        ttk.Label(options_frame, text="输出文件名:").grid(row=1, column=0, padx=(0, 10), pady=(5, 0), sticky=tk.W)
        self.filename_entry = ttk.Entry(options_frame, font=("Arial", 10))
        self.filename_entry.grid(row=1, column=1, columnspan=2, pady=(5, 0), sticky=(tk.W, tk.E)) # columnspan=2 让输入框填满
    
        ttk.Label(options_frame, text="转换格式:").grid(row=2, column=0, pady=(10, 0), padx=(0, 10), sticky=tk.W)
        self.format_var = tk.StringVar(value="docx")
        self.format_combo = ttk.Combobox(options_frame, textvariable=self.format_var,
                                       values=["docx", "txt"], state="readonly", width=15)
        self.format_combo.grid(row=2, column=1, pady=(10, 0), columnspan=2, sticky=tk.W)
        self.format_combo.bind("<<ComboboxSelected>>", self.update_filename)
    
        # ---- 转换按钮和状态栏 ----
        action_frame = ttk.Frame(main_frame)
        action_frame.grid(row=3, column=0, pady=(20, 0), sticky=(tk.W, tk.E))
        action_frame.columnconfigure(0, weight=1)
    
        self.convert_btn = ttk.Button(action_frame, text="开始转换", command=self.convert_file, style="Accent.TButton")
        self.convert_btn.grid(row=0, column=0, ipady=5, sticky=(tk.W, tk.E))
        
        self.status_label = ttk.Label(action_frame, text="准备就绪", anchor=tk.W, foreground="gray")
        self.status_label.grid(row=1, column=0, pady=(10, 0), sticky=(tk.W, tk.E))
        self.update_status("准备就绪", "gray")
    
    def browse_html(self):
        filepath = filedialog.askopenfilename(
            title="选择HTML文件",
            filetypes=[("HTML文件", "*.html"), ("所有文件", "*.*")]
        )
        if filepath:
            self.html_entry.delete(0, tk.END)
            self.html_entry.insert(0, filepath)
            
            # --- 新增: 自动设置默认输出目录 ---
            output_dir = os.path.dirname(filepath)
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, output_dir)
            
            self.update_filename()
            self.update_status(f"已选择文件: {os.path.basename(filepath)}", "blue")
            
    # --- 新增: 选择输出目录的方法 ---
    def browse_output_dir(self):
        dir_path = filedialog.askdirectory(title="选择输出目录")
        if dir_path:
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, dir_path)
            self.update_status(f"输出目录已设置为: {dir_path}", "blue")

    def update_filename(self, event=None):
        html_path = self.html_entry.get()
        if html_path and os.path.isfile(html_path):
            base_name = os.path.splitext(os.path.basename(html_path))[0]
            format_ext = self.format_var.get()
            default_name = f"{base_name}.{format_ext}"
            self.filename_entry.delete(0, tk.END)
            self.filename_entry.insert(0, default_name)
    
    def convert_file(self):
        html_path = self.html_entry.get()
        # --- 新增: 获取输出目录 ---
        output_dir = self.output_dir_entry.get()
        output_filename = self.filename_entry.get()
        format_type = self.format_var.get()
    
        if not html_path or not os.path.isfile(html_path):
            messagebox.showerror("错误", "请选择有效的HTML文件")
            return
            
        # --- 新增: 验证输出目录 ---
        if not output_dir or not os.path.isdir(output_dir):
            messagebox.showerror("错误", "请选择一个有效的输出目录")
            return
            
        if not output_filename:
            messagebox.showerror("错误", "请输入输出文件名")
            return
            
        self.convert_btn.config(state=tk.DISABLED)
        self.update_status("正在转换中，请稍候...", "orange")
        self.root.update_idletasks() # 强制更新UI
    
        try:
            # --- 修改: 使用指定的输出目录构建路径 ---
            output_path = os.path.join(output_dir, output_filename)
            
            if format_type == "txt":
                self.convert_to_txt(html_path, output_path)
            elif format_type == "docx":
                self.convert_to_docx(html_path, output_path)
            
            messagebox.showinfo("成功", f"文件已成功转换为 {format_type.upper()} 格式！\n\n路径: {output_path}")
            self.update_status("转换完成！", "green")
    
        except Exception as e:
            messagebox.showerror("错误", f"转换失败: {str(e)}")
            self.update_status(f"转换失败: {e}", "red")
        finally:
            self.convert_btn.config(state=tk.NORMAL)
    
    def update_status(self, text, color):
        """更新状态标签的文本和颜色"""
        self.status_label.config(text=text, foreground=color)
    
    def convert_to_txt(self, html_path, output_path):
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    def convert_to_docx(self, html_path, output_path):
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        doc = docx.Document()
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
            if element.name.startswith('h'):
                try:
                    heading_level = int(element.name[1])
                    level = min(heading_level, 4)
                except (ValueError, IndexError):
                    level = 1
                paragraph = doc.add_heading('', level=level)
                self.add_text_to_paragraph(paragraph, element.get_text())
            else:
                paragraph = doc.add_paragraph()
                self.add_text_to_paragraph(paragraph, element.get_text())
        doc.save(output_path)
    
    def add_text_to_paragraph(self, paragraph, text):
        run = paragraph.add_run(text)
        run.font.size = Pt(12)

if __name__ == "__main__":
    root = tk.Tk()
    app = HTMLConverter(root)
    root.mainloop()